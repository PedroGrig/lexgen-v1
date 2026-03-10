"""
generator.py — Gera arquivos .docx formatados a partir do texto gerado pelo LLM.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx(content: str, document_type_label: str) -> bytes:
    """
    Cria um arquivo .docx formatado com o conteúdo gerado.

    - Fonte: Arial 12
    - Espaçamento: 1.5
    - Margens: padrão ABNT (3cm superior/esquerda, 2cm inferior/direita)
    - Cabeçalho com nome do documento
    """
    doc = Document()

    # Configurar margens ABNT
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Configurar cabeçalho
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(document_type_label.upper())
    header_run.font.name = "Arial"
    header_run.font.size = Pt(10)
    header_run.font.bold = True

    # Estilo padrão do documento
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    # Processar o conteúdo linha por linha
    lines = content.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Linha em branco — adiciona parágrafo vazio
            doc.add_paragraph("")
            continue

        # Detectar títulos/seções (linhas em caixa alta ou começando com números romanos)
        is_title = (
            stripped.isupper()
            and len(stripped) > 3
            and len(stripped) < 100
        )

        # Detectar subtítulos (começam com números, alíneas ou marcadores)
        is_subtitle = (
            stripped.startswith(("I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII.", "IX.", "X."))
            or stripped.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."))
            or stripped.startswith(("a)", "b)", "c)", "d)", "e)", "f)"))
        )

        para = doc.add_paragraph()

        if is_title:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        elif is_subtitle:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            para.paragraph_format.space_before = Pt(6)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(12)

        # Indentação de primeira linha para parágrafos normais
        if not is_title and not is_subtitle and len(stripped) > 50:
            para.paragraph_format.first_line_indent = Cm(1.25)

    # Salvar em bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
